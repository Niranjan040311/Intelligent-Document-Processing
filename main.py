import asyncio
import hashlib
import io
import json
import logging
import os
import re
import threading
import time
from contextlib import asynccontextmanager
from pathlib import Path
from typing import List, Optional

import cv2
import easyocr
import fitz  # PyMuPDF
import numpy as np
from docx import Document
import openpyxl
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from PIL import Image
from pydantic import BaseModel

# Load a local .env file (if present) so API keys can live there instead of
# being typed into the terminal each run. .env is git-ignored — never commit it.
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv optional; OS env vars still work without it

# ── Logging ────────────────────────────────────────────────────────────────
# Step-by-step diagnostics so a failing document can be pinpointed in the
# server console. Tune verbosity with the level below (DEBUG → INFO → WARNING).
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-7s | %(name)s | %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("docscan")

UPLOAD_DIR = Path("uploads")
OCR_DIR = Path("ocr_results")
UPLOAD_DIR.mkdir(exist_ok=True)
OCR_DIR.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg", ".txt", ".xlsx", ".xls"}
MAX_FILE_SIZE_MB = 20

# ── LLM fallback (GPT-4.1) ───────────────────────────────────────────────────
# When keyword-proximity can't find a field, that field's name + the document
# text are sent to GPT-4.1, which understands meaning ("MR. Ravi is the patient"
# → name = "MR. Ravi") even with no label present.
#
# Two hosting modes, auto-detected from environment variables:
#   • Azure OpenAI  — set AZURE_OPENAI_ENDPOINT + AZURE_OPENAI_API_KEY
#                     (+ AZURE_OPENAI_DEPLOYMENT, AZURE_OPENAI_API_VERSION)
#   • OpenAI direct — set OPENAI_API_KEY (+ optional OPENAI_MODEL)
# If neither is configured, the app stays fully offline (keyword extraction only).
#
# Keys are read from the environment / .env — NEVER hardcode them in source.
AZURE_OPENAI_ENDPOINT    = os.environ.get("AZURE_OPENAI_ENDPOINT", "").strip()
AZURE_OPENAI_API_KEY     = os.environ.get("AZURE_OPENAI_API_KEY", "").strip()
AZURE_OPENAI_DEPLOYMENT  = os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4.1").strip()
AZURE_OPENAI_API_VERSION = os.environ.get("AZURE_OPENAI_API_VERSION", "2024-10-21").strip()

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "").strip()
OPENAI_MODEL   = os.environ.get("OPENAI_MODEL", "gpt-4.1").strip()

_llm_client = None    # lazy-initialised on first fallback call
_llm_mode = None      # "azure" | "openai" | None

ocr_reader: easyocr.Reader = None
# EasyOCR's Reader is NOT thread-safe — concurrent readtext() calls on the same
# instance can crash (which surfaced as random "FAILED" docs in batches). This
# lock serialises OCR so only one readtext runs at a time.
_ocr_lock = threading.Lock()

# In-memory OCR cache: MD5(file bytes) → extracted text
# Same file uploaded again returns instantly without re-running OCR.
_ocr_cache: dict[str, str] = {}

# Numeric field keywords — for these, prefer the LAST high-confidence match
# because grand totals appear at the bottom of invoices, not the top.
_NUMERIC_KEYWORDS = {
    "total", "amount", "cgst", "sgst", "igst", "tax", "price",
    "cost", "fee", "charge", "balance", "due", "payable", "subtotal",
    "discount", "gross", "net",
}


@asynccontextmanager
async def lifespan(app):
    global ocr_reader
    print("Loading EasyOCR model...")
    # EasyOCR model loading is CPU/IO bound — run in thread so startup doesn't block.
    ocr_reader = await asyncio.to_thread(easyocr.Reader, ["en"], gpu=False)
    print("EasyOCR ready.")
    yield


class ExtractRequest(BaseModel):
    text: str
    fields: List[str]


class LLMExtractRequest(BaseModel):
    text: str
    fields: List[str]   # only the fields keyword-proximity could not find


app = FastAPI(title="DocScan IDP API", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


def is_allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def _md5(data: bytes) -> str:
    return hashlib.md5(data, usedforsecurity=False).hexdigest()


# ── Image preprocessing ──────────────────────────────────────────────────────

def preprocess_for_ocr(img_array: np.ndarray) -> np.ndarray:
    gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY) if img_array.ndim == 3 else img_array.copy()

    h, w = gray.shape
    if h < 1200 or w < 900:
        scale = max(1200 / h, 900 / w)
        gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

    gray = cv2.fastNlMeansDenoising(gray, h=8, templateWindowSize=7, searchWindowSize=21)

    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    gray = clahe.apply(gray)

    blurred = cv2.GaussianBlur(gray, (0, 0), 3)
    gray = cv2.addWeighted(gray, 1.5, blurred, -0.5, 0)
    gray = np.clip(gray, 0, 255).astype(np.uint8)

    return gray


# ── OCR text post-processing ─────────────────────────────────────────────────

def postprocess_ocr_text(text: str) -> str:
    text = re.sub(r'[<&]\s*(?=\d)', '₹', text)
    text = re.sub(r'\b8(\d{1,2},\d{3})', r'₹\1', text)
    text = re.sub(r'\b(Rs\.?|INR)\s*', '₹', text, flags=re.IGNORECASE)
    text = re.sub(r'(?<=\d)[Oo](?=\d)', '0', text)
    text = re.sub(r'(?<=\d)l(?=\d)', '1', text)
    text = re.sub(r'\bl(?=\d{2,})', '1', text)
    text = re.sub(r'(@[\w-]+)\s+(com|in|org|net|io|co)\b', r'\1.\2', text, flags=re.IGNORECASE)
    text = re.sub(r'\b([\w-]{3,})\s+(com|in|org|net|io)\b(?![\w.])', r'\1.\2', text, flags=re.IGNORECASE)
    text = re.sub(r'  +', ' ', text)
    return text


# ── Core OCR ─────────────────────────────────────────────────────────────────

def ocr_image_array(img_array: np.ndarray) -> str:
    processed = preprocess_for_ocr(img_array)
    # Serialise access — the shared reader can't run two readtext() at once.
    with _ocr_lock:
        results = ocr_reader.readtext(processed)
    raw = "\n".join(text for _, text, conf in results if conf > 0.3)
    return postprocess_ocr_text(raw)


def _extract_text_sync(file_bytes: bytes, filename: str) -> str:
    """CPU-bound extraction — always called via asyncio.to_thread."""
    ext = Path(filename).suffix.lower()
    log.info("OCR start: %s (%s, %d bytes)", filename, ext, len(file_bytes))

    if ext == ".pdf":
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        log.info("PDF '%s' has %d page(s)", filename, doc.page_count)
        pages = []
        for i, page in enumerate(doc):
            try:
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                text = ocr_image_array(np.array(img))
                log.info("PDF '%s' page %d/%d OCR'd → %d chars",
                         filename, i + 1, doc.page_count, len(text))
                pages.append(f"--- Page {i + 1} ---\n{text}")
            except Exception:
                log.exception("PDF '%s' page %d FAILED during OCR", filename, i + 1)
                pages.append(f"--- Page {i + 1} ---\n[OCR failed for this page]")
        return "\n\n".join(pages)

    if ext in {".png", ".jpg", ".jpeg"}:
        img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        return ocr_image_array(np.array(img))

    if ext in {".docx", ".doc"}:
        doc = Document(io.BytesIO(file_bytes))
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return postprocess_ocr_text("\n".join(parts))

    if ext in {".xlsx", ".xls"}:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return postprocess_ocr_text("\n".join(parts))

    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")

    return ""


async def extract_text(file_bytes: bytes, filename: str) -> str:
    """Async wrapper with MD5 cache — same file never re-OCR'd."""
    key = _md5(file_bytes)
    if key in _ocr_cache:
        log.info("Cache HIT for '%s' (md5=%s)", filename, key[:8])
        return _ocr_cache[key]

    log.info("Cache MISS for '%s' (md5=%s) — running OCR", filename, key[:8])
    started = time.perf_counter()
    try:
        result = await asyncio.to_thread(_extract_text_sync, file_bytes, filename)
    except Exception:
        log.exception("OCR FAILED for '%s'", filename)
        raise
    elapsed = time.perf_counter() - started
    log.info("OCR done: '%s' → %d chars in %.2fs", filename, len(result), elapsed)
    if not result.strip():
        log.warning("OCR for '%s' produced EMPTY text — possible blank/unreadable document", filename)
    _ocr_cache[key] = result
    return result


def save_ocr_result(filename: str, text: str) -> Path:
    out_path = OCR_DIR / f"{Path(filename).stem}.txt"
    out_path.write_text(text, encoding="utf-8")
    return out_path


# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/", include_in_schema=False)
def root():
    return FileResponse("static/index.html")


app.mount("/static", StaticFiles(directory="static"), name="static")


@app.post("/upload", summary="Upload & OCR a document")
async def upload_file(file: UploadFile = File(...)):
    log.info("POST /upload received: '%s'", file.filename)
    if not is_allowed_file(file.filename):
        log.warning("Rejected '%s' — file type not allowed", file.filename)
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed. Accepted: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    contents = await file.read()
    if len(contents) > MAX_FILE_SIZE_MB * 1024 * 1024:
        log.warning("Rejected '%s' — %d bytes exceeds %d MB limit",
                    file.filename, len(contents), MAX_FILE_SIZE_MB)
        raise HTTPException(status_code=400, detail=f"File exceeds {MAX_FILE_SIZE_MB} MB limit")

    (UPLOAD_DIR / file.filename).write_bytes(contents)

    try:
        # Non-blocking: OCR runs in a thread pool, event loop stays free
        extracted_text = await extract_text(contents, file.filename)
    except Exception as e:
        log.exception("POST /upload FAILED for '%s'", file.filename)
        raise HTTPException(status_code=500, detail=f"OCR failed: {e}")

    ocr_path = save_ocr_result(file.filename, extracted_text)
    log.info("POST /upload OK: '%s' → saved %s", file.filename, ocr_path)

    return JSONResponse(status_code=200, content={
        "filename": file.filename,
        "size_bytes": len(contents),
        "ocr_saved_to": str(ocr_path),
        "extracted_text": extracted_text,
    })


@app.post("/upload/multiple", summary="Upload & OCR multiple documents")
async def upload_multiple(files: List[UploadFile] = File(...)):
    async def _process(file: UploadFile) -> dict:
        if not is_allowed_file(file.filename):
            return {"filename": file.filename, "status": "rejected", "reason": "File type not allowed"}

        contents = await file.read()
        if len(contents) > MAX_FILE_SIZE_MB * 1024 * 1024:
            return {"filename": file.filename, "status": "rejected", "reason": "File too large"}

        (UPLOAD_DIR / file.filename).write_bytes(contents)
        extracted_text = await extract_text(contents, file.filename)
        ocr_path = save_ocr_result(file.filename, extracted_text)

        return {
            "filename": file.filename,
            "status": "processed",
            "size_bytes": len(contents),
            "ocr_saved_to": str(ocr_path),
            "extracted_text": extracted_text,
        }

    # All files processed concurrently — not sequentially
    results = await asyncio.gather(*[_process(f) for f in files])
    return JSONResponse(status_code=200, content={"results": list(results)})


@app.get("/files", summary="List uploaded files and their OCR results")
def list_files():
    files = [
        {
            "filename": f.name,
            "size_bytes": f.stat().st_size,
            "ocr_result": str(OCR_DIR / f"{f.stem}.txt") if (OCR_DIR / f"{f.stem}.txt").exists() else None,
        }
        for f in UPLOAD_DIR.iterdir() if f.is_file()
    ]
    return {"files": files, "count": len(files)}


# ── LLM fallback extraction (GPT-4.1) ────────────────────────────────────────

def _llm_available() -> bool:
    """True if either Azure OpenAI or OpenAI-direct is configured."""
    return bool(AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_API_KEY) or bool(OPENAI_API_KEY)


def _load_llm_client():
    """Lazy-init the OpenAI client (Azure or direct), first call only."""
    global _llm_client, _llm_mode
    if _llm_client is not None:
        return _llm_client

    if AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_API_KEY:
        from openai import AzureOpenAI
        _llm_client = AzureOpenAI(
            azure_endpoint=AZURE_OPENAI_ENDPOINT,
            api_key=AZURE_OPENAI_API_KEY,
            api_version=AZURE_OPENAI_API_VERSION,
        )
        _llm_mode = "azure"
        log.info("LLM fallback ready (Azure OpenAI, deployment=%s)", AZURE_OPENAI_DEPLOYMENT)
    elif OPENAI_API_KEY:
        from openai import OpenAI
        _llm_client = OpenAI(api_key=OPENAI_API_KEY)
        _llm_mode = "openai"
        log.info("LLM fallback ready (OpenAI direct, model=%s)", OPENAI_MODEL)

    return _llm_client


def _llm_extract_fields(text: str, fields: list[str]) -> Optional[dict[str, dict]]:
    """
    Ask GPT-4.1 to extract the given fields from `text` in ONE call.

    Returns {field: {"value": str, "confidence": float}} or None on any failure
    (caller then leaves those fields as "Not detected").
    """
    if not _llm_available() or not fields:
        return None

    try:
        client = _load_llm_client()
        model = AZURE_OPENAI_DEPLOYMENT if _llm_mode == "azure" else OPENAI_MODEL

        field_list = "\n".join(f"  - {f}" for f in fields)
        system_msg = (
            "You are a meticulous document data-extraction engine. Your only job is to "
            "find the value for each requested field inside raw OCR'd document text and "
            "return it as strict JSON. Accuracy matters far more than completeness — a "
            'wrong value is worse than "Not detected".\n'
            "\n"
            "EXTRACTION RULES:\n"
            "1. Use meaning and context, not just labels. The value may appear with a "
            'label ("Patient: John Doe"), in prose ("Mr. Ravi is the patient"), in a '
            "table, or implied by position. Infer which span of text answers the field.\n"
            "2. Return ONLY the value itself — no field name, no label, no quotes, no "
            'units unless the unit IS the value, no trailing punctuation. For "patient '
            'name" return "John Doe", never "Patient: John Doe".\n'
            "3. The OCR text may contain errors (0/O, 1/l/I, missing spaces, broken "
            "lines, merged words). Correct OBVIOUS OCR errors in the value when you are "
            "confident of the intended text; otherwise return it as-is. Never invent "
            "characters that are not implied by the text.\n"
            "4. NEVER hallucinate. If the document does not actually contain the value, "
            'you MUST return "Not detected". Do not guess plausible-but-absent values '
            "(e.g. do not invent a date or an amount that is not written in the text).\n"
            "5. If several candidates exist, pick the one that best matches the field's "
            "intent. For totals/grand totals prefer the final/largest summary figure; "
            "for names prefer the primary subject of the document.\n"
            "6. Normalise lightly: trim surrounding whitespace; keep the value's own "
            "formatting (dates, currency symbols, IDs) as written in the document.\n"
            "\n"
            "CONFIDENCE (0-100) — be calibrated, not optimistic:\n"
            "  90-100 = value is explicitly and unambiguously stated.\n"
            "  70-89  = clearly implied or has minor OCR cleanup.\n"
            "  40-69  = inferred from context, or more than one candidate existed.\n"
            "  1-39   = weak guess; the document is ambiguous.\n"
            '  0      = not present — value MUST be exactly "Not detected".\n'
            "\n"
            "OUTPUT: a single JSON object. Keys EXACTLY match the requested field names. "
            'Each maps to {"value": <string>, "confidence": <number 0-100>}. No prose, '
            "no markdown, no extra keys."
        )
        user_msg = (
            "Extract these fields. Return JSON whose keys are EXACTLY these names "
            '(case and spelling unchanged), each mapping to {"value", "confidence"}:\n'
            f"{field_list}\n"
            "\n"
            "EXAMPLE (for shape only — use the real document below for actual values):\n"
            '  fields: ["patient name", "amount due", "email"]\n'
            '  text:   "Mr. Ravi is the patient. Total payable Rs.4,500."\n'
            '  output: {"patient name": {"value": "Ravi", "confidence": 88}, '
            '"amount due": {"value": "Rs.4,500", "confidence": 95}, '
            '"email": {"value": "Not detected", "confidence": 0}}\n'
            "\n"
            f"--- DOCUMENT TEXT ---\n{text}\n--- END DOCUMENT TEXT ---"
        )

        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg},
            ],
            temperature=0,
            response_format={"type": "json_object"},
        )
        parsed = json.loads(resp.choices[0].message.content)
        log.info("LLM fallback extracted %d field(s) via %s", len(parsed), _llm_mode)
        return parsed

    except Exception:
        log.exception("LLM fallback failed — leaving fields as not-detected")
        return None


# ── Field extraction ──────────────────────────────────────────────────────────

def _is_numeric_field(field_lower: str) -> bool:
    return any(kw in field_lower for kw in _NUMERIC_KEYWORDS)


def _whole_word(haystack_lower: str, needle_lower: str) -> bool:
    """Match `needle` only as a complete word/phrase, not inside another word.

    Prevents a short field like "id" from matching inside "paid"/"void", and
    keeps "Order ID" from being mistaken for the field "ID".
    """
    if not needle_lower:
        return False
    return re.search(rf'(?<!\w){re.escape(needle_lower)}(?!\w)', haystack_lower) is not None


def extract_single_field(lines: list[str], field: str) -> dict:
    field_lower = field.lower().strip()
    # Significant words only (drop tiny connector words), but keep the field
    # usable even when it's a single short token like "id" or "po".
    field_words = [w for w in re.findall(r'\w+', field_lower) if len(w) > 2]

    candidates: list[tuple[float, int, str]] = []  # (confidence, line_index, value)

    for i, line in enumerate(lines):
        line_lower = line.lower()

        # The label portion is what sits before a ':'/'=' — the field name should
        # really live there, not buried in the value.
        split = re.split(r'[:=]\s*', line, maxsplit=1)
        label_lower = split[0].lower() if len(split) > 1 else line_lower

        if _whole_word(line_lower, field_lower):
            # Full phrase present as whole words. Strongest when it's the label.
            conf_base = 0.95 if _whole_word(label_lower, field_lower) else 0.90
            # Exact label ("ID: 4471" for field "ID") beats a prefixed label
            # ("Customer ID: ...") so we stop grabbing the first loose match.
            if label_lower.strip() == field_lower:
                conf_base = 0.98
        elif field_words and all(_whole_word(line_lower, w) for w in field_words):
            conf_base = 0.80
        elif field_words and len(field_words) > 1:
            hit_ratio = sum(1 for w in field_words if _whole_word(line_lower, w)) / len(field_words)
            if hit_ratio >= 0.6:
                conf_base = hit_ratio * 0.68
            else:
                continue
        else:
            continue

        same_line_split = split
        if len(same_line_split) > 1 and same_line_split[1].strip():
            candidate = same_line_split[1].strip()
            conf = conf_base
        else:
            next_idx = i + 1
            while next_idx < len(lines) and not lines[next_idx].strip():
                next_idx += 1
            if next_idx < len(lines) and len(lines[next_idx]) < 200:
                candidate = lines[next_idx].strip()
                conf = conf_base * 0.87
            else:
                continue

        candidates.append((conf, i, candidate))

    if not candidates:
        return {"field": field, "value": "Not detected", "confidence": 0.0}

    if _is_numeric_field(field_lower):
        # Numeric fields: among candidates within 15% of best confidence,
        # prefer the LAST occurrence — grand totals sit at the bottom of documents.
        best_conf = max(c[0] for c in candidates)
        threshold = best_conf * 0.85
        eligible = [c for c in candidates if c[0] >= threshold]
        conf, _, value = eligible[-1]
    else:
        # Non-numeric fields: highest confidence wins
        conf, _, value = max(candidates, key=lambda x: x[0])

    return {"field": field, "value": value, "confidence": round(min(conf * 100, 99.0), 1)}


# ── Field validation ──────────────────────────────────────────────────────────
# Sanity-check each extracted value so low-confidence or malformed data is
# flagged for review instead of silently trusted.

# Found, well-formed values below this confidence are flagged "review".
REVIEW_CONFIDENCE = 60.0

# Keyword matches below this confidence are re-checked by GPT-4.1 (Tier 2),
# since a low-confidence grab is often the wrong span of text.
LLM_RECHECK_CONFIDENCE = 90.0

_DATE_RE = re.compile(
    r'\b('
    r'\d{1,2}[/\-.]\d{1,2}[/\-.]\d{2,4}'        # 12/05/2024, 12-05-24
    r'|\d{4}[/\-.]\d{1,2}[/\-.]\d{1,2}'          # 2024-05-12
    r'|\d{1,2}\s+[A-Za-z]{3,9},?\s+\d{2,4}'      # 12 May 2024
    r'|[A-Za-z]{3,9}\s+\d{1,2},?\s+\d{2,4}'      # May 12, 2024
    r')\b'
)
_EMAIL_RE = re.compile(r'[^@\s]+@[^@\s]+\.[^@\s]+')
_NUMBER_RE = re.compile(r'\d')


def _field_kind(field_lower: str) -> str:
    """Guess the expected data type from the field name."""
    if any(k in field_lower for k in ("email", "e-mail")) or field_lower.endswith("mail"):
        return "email"
    if any(k in field_lower for k in ("date", "dob", "dated", "expiry", "issued", "due on")):
        return "date"
    if any(k in field_lower for k in ("phone", "mobile", "contact", "tel", "cell")):
        return "phone"
    if _is_numeric_field(field_lower):
        return "amount"
    return "text"


def validate_field(field: str, value: str, confidence: float) -> dict:
    """Validate one extracted value. Returns {status, message}.

    status ∈ {valid, review, invalid, missing}:
      missing  — nothing was extracted
      invalid  — value doesn't match the expected format for the field
      review   — well-formed but low confidence; a human should confirm
      valid    — found, well-formed, high confidence
    """
    if value in ("Not detected", "Error") or not value.strip():
        return {"status": "missing", "message": "Not found in document"}

    kind = _field_kind(field.lower())
    v = value.strip()

    if kind == "email" and not _EMAIL_RE.search(v):
        return {"status": "invalid", "message": "Doesn't look like an email address"}
    if kind == "date" and not _DATE_RE.search(v):
        return {"status": "invalid", "message": "Doesn't look like a date"}
    if kind == "phone":
        digits = re.sub(r'\D', '', v)
        if not 7 <= len(digits) <= 15:
            return {"status": "invalid", "message": "Doesn't look like a phone number"}
    if kind == "amount" and not _NUMBER_RE.search(v):
        return {"status": "invalid", "message": "Expected a numeric value"}

    if confidence < REVIEW_CONFIDENCE:
        return {"status": "review", "message": f"Low confidence ({confidence:.0f}%) — please verify"}

    return {"status": "valid", "message": "Looks valid"}


@app.post("/extract", summary="Extract specific fields from OCR text with confidence scores")
async def extract_fields(req: ExtractRequest):
    log.info("POST /extract: %d field(s) over %d chars — fields=%s",
             len(req.fields), len(req.text), req.fields)
    if not req.text.strip():
        log.warning("POST /extract rejected — empty text")
        raise HTTPException(status_code=400, detail="text cannot be empty")
    if not req.fields:
        log.warning("POST /extract rejected — empty fields list")
        raise HTTPException(status_code=400, detail="fields list cannot be empty")

    lines = [line.strip() for line in req.text.splitlines()]

    # ── Tier 1 only: keyword-proximity (offline, fast). The GPT-4.1 fallback
    # is a SEPARATE endpoint (/extract/llm) so the UI can show the two tiers
    # firing in sequence. This call returns instantly with the offline result. ──
    try:
        results = await asyncio.gather(
            *[asyncio.to_thread(extract_single_field, lines, f) for f in req.fields]
        )
    except Exception as e:
        log.exception("POST /extract FAILED")
        raise HTTPException(status_code=500, detail=f"Extraction failed: {e}")
    results = list(results)
    for r in results:
        r.setdefault("source", "keyword")
        r["validation"] = validate_field(r["field"], r["value"], r["confidence"])

    found = sum(1 for r in results if r["value"] != "Not detected")
    needs_review = sum(1 for r in results if r["validation"]["status"] in ("review", "invalid"))
    overall = round(sum(r["confidence"] for r in results) / len(results), 1) if results else 0

    # Fields worth a Tier-2 (GPT-4.1) pass: not just MISSING, but also
    # SUSPICIOUS keyword matches — low confidence or failed validation. This
    # catches confident-but-wrong grabs (e.g. field "Name" matching a table
    # header and returning "Primary?"). GPT-4.1's answer can then override.
    recheck = [
        r["field"] for r in results
        if r["value"] == "Not detected"
        or r["confidence"] < LLM_RECHECK_CONFIDENCE
        or r["validation"]["status"] in ("invalid", "review")
    ]
    log.info("POST /extract OK [tier1]: %d/%d found, %d to recheck, %d need review, accuracy %.1f%%",
             found, len(results), len(recheck), needs_review, overall)

    return {
        "extractions": results,
        "overall_accuracy": overall,
        "fields_found": found,
        "fields_total": len(results),
        "needs_review": needs_review,
        # Fields the UI should send to Tier-2 GPT-4.1 (missing OR suspicious).
        "recheck_fields": recheck,
        "llm_available": _llm_available(),
    }


@app.post("/extract/llm", summary="Tier 2: recover unfound fields via GPT-4.1")
async def extract_fields_llm(req: LLMExtractRequest):
    """Second tier of the extraction flow. The UI calls this with ONLY the
    fields that keyword-proximity could not find, so it can show 'Asking
    GPT-4.1 for N fields…' as a distinct live step."""
    log.info("POST /extract/llm: %d field(s) — fields=%s", len(req.fields), req.fields)
    if not _llm_available():
        return {"extractions": [], "llm_used": False, "reason": "LLM not configured"}
    if not req.text.strip() or not req.fields:
        return {"extractions": [], "llm_used": False, "reason": "nothing to do"}

    llm_data = await asyncio.to_thread(_llm_extract_fields, req.text, req.fields)
    if not llm_data:
        log.info("POST /extract/llm: GPT-4.1 returned nothing")
        return {"extractions": [], "llm_used": False, "reason": "no response"}

    out = []
    for field in req.fields:
        entry = llm_data.get(field) or {}
        value = str(entry.get("value", "")).strip() or "Not detected"
        if value == "Not detected":
            continue
        try:
            conf = round(min(float(entry.get("confidence", 0)), 99.0), 1)
        except (TypeError, ValueError):
            conf = 0.0
        r = {"field": field, "value": value, "confidence": conf, "source": "llm"}
        r["validation"] = validate_field(field, value, conf)
        out.append(r)

    log.info("POST /extract/llm OK: GPT-4.1 recovered %d/%d field(s)", len(out), len(req.fields))
    return {"extractions": out, "llm_used": len(out) > 0}


@app.get("/health")
def health():
    llm_mode = (
        "azure" if (AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_API_KEY)
        else "openai" if OPENAI_API_KEY
        else None
    )
    return {
        "status": "ok",
        "ocr_ready": ocr_reader is not None,
        "cache_entries": len(_ocr_cache),
        "llm_fallback_enabled": _llm_available(),
        "llm_mode": llm_mode,
    }
